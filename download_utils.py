from fpdf import FPDF
from docx import Document
import os

def export_to_pdf(chats, filename):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    for chat in chats:
        pdf.multi_cell(0, 10, f"Q: {chat['question']}")
        pdf.multi_cell(0, 10, f"A: {chat['answer']}\n")

    pdf.output(filename)

def export_to_docx(chats, filename):
    doc = Document()
    doc.add_heading("Chat History", 0)

    for chat in chats:
        doc.add_paragraph(f"Q: {chat['question']}", style='List Bullet')
        doc.add_paragraph(f"A: {chat['answer']}\n")

    doc.save(filename)

def delete_file_if_exists(filepath):
    if os.path.exists(filepath):
        os.remove(filepath)
