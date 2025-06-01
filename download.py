from fpdf import FPDF
from docx import Document
import os

def export_to_pdf(chats, filename="user_chats.pdf"):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    for chat in chats:
        pdf.multi_cell(0, 10, f"Q: {chat['question']}")
        pdf.ln(2)
        pdf.multi_cell(0, 10, f"A: {chat['answer']}")
        pdf.ln(10)

    pdf.output(filename)
    return filename

def export_to_docx(chats, filename="user_chats.docx"):
    doc = Document()
    doc.add_heading("User Chat History", 0)

    for chat in chats:
        doc.add_paragraph("Q: " + chat['question'], style='List Bullet')
        doc.add_paragraph("A: " + chat['answer'], style='List Paragraph')
        doc.add_paragraph("")  # Add space

    doc.save(filename)
    return filename

def delete_file_if_exists(filename):
    try:
        if os.path.exists(filename):
            os.remove(filename)
    except Exception as e:
        print(f"Error deleting file {filename}: {str(e)}")
